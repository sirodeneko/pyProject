import json
from datetime import timedelta

from django.db import transaction
from django.http import HttpResponse, StreamingHttpResponse
from django.shortcuts import render

from django.utils.http import urlquote
from docx.text.paragraph import Paragraph
from docx.styles import style
from docx import Document
from docx.document import Document as Doc
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from project.web.entity.areaInfo import models
from project.web.utils import updata, create_image, email


def index(request):
    versions = models.Versions.objects.last()
    # 如果没有数据 就去拿数据
    if versions is None:
        return fetch(request)

    area_infos = models.AreaInfo.objects.filter(v_id=versions).order_by("-max_num")

    # for item in area_infos:
    #     print(item.name)

    context = {
        'data': area_infos,
        'i': 0,
    }
    return render(request, 'info_index.html', context)


def fetch(request):
    data = updata.getData()
    # 开启事务
    with transaction.atomic():
        version = models.Versions.objects.create()
        list_to_insert = []
        for row in data["Rows"]:
            item = models.AreaInfo(
                code=row["CODE"],
                name=row["NAME"],
                address=row["ADDRESS"],
                des=row["DES"],
                time=row["TIME"],
                grade=row["GRADE"],
                t_time=row["T_TIME"],
                max_num=row["MAX_NUM"] if row["MAX_NUM"] != "" else 0,
                ssd=row["SSD"],
                num=row["NUM"] if row["NUM"] != "" else 0,
                type=row["TYPE"],
                t_code=row["T_CODE"],
                v_id=version,
            )
            list_to_insert.append(item)
        # 批量插入
        models.AreaInfo.objects.bulk_create(list_to_insert)

    # 查询数据
    area_infos = models.AreaInfo.objects.filter(v_id=version).order_by("-max_num")
    # models.Versions.delete(version)
    context = {
        'data': area_infos,
        'i': 0,
    }

    return render(request, 'info_index.html', context)


def get_area_list():
    versions = models.Versions.objects.last()
    area_infos = models.AreaInfo.objects.filter(v_id=versions).order_by("-max_num")
    return area_infos


def get_single_chart(request):
    # 获取传过来的code 编号
    t_code = request.GET.get('code', default='-1')
    # id逆序取前7个
    v_list = models.Versions.objects.order_by("-id").all()[:7]
    vv_list = []
    vt_list = []
    # 获取vv_list(拿到版本号，用于之后in查询）
    # 获取vt_list(时间列表）
    for item in v_list:
        vv_list.append(item.id)
        # utc时区格式化为东八区（+8h)
        vt_list.append((item.v_time + timedelta(hours=8)).strftime("%m-%d %H"))
    # 反转（因为是逆序查询出来的）
    vt_list.reverse()
    # in查询出数据
    area_infos = models.AreaInfo.objects.filter(v_id__in=vv_list, code=t_code).all()
    context = {
        'name_list': get_area_list(),  # 名字列表
        'data': area_infos,  # 主要的数据
        'tName': area_infos[0].name,  # 获取当前的名字
        'tTime': vt_list,  # 时间列表
        'i': 0,
    }
    return render(request, 'singleChart.html', context)


def get_multi_chart(request):
    # 拿到参数列表
    t_code = request.GET.getlist('code', default=[])
    # 判断是否为空
    if t_code is []:
        context = {
            'name_list': get_area_list(),
        }
        return render(request, 'multiChart.html', context)

    # 老规矩
    v_list = models.Versions.objects.order_by("-id").all()[:7]
    vv_list = []
    vt_list = []
    for item in v_list:
        vv_list.append(item.id)
        # utc时区格式化为东八区（+8h)
        vt_list.append((item.v_time + timedelta(hours=8)).strftime("%m-%d %H"))

    vt_list.reverse()
    areas_infos = []
    # 对于每个参数code都进行一次查询 丢到areas_infos中去
    for i_code in t_code:
        areas_infos.append(models.AreaInfo.objects.filter(v_id__in=vv_list, code=i_code).all())

    # 获取每个对应的名字，因为在模板中不方便获取这个值
    t_name_list = []
    for item in areas_infos:
        t_name_list.append(item.first().name)
    context = {
        'name_list': get_area_list(),
        't_name_list': t_name_list,
        'data': areas_infos,
        'tTime': vt_list,
    }
    return render(request, 'multiChart.html', context)


def new_docx_report(request):
    # 拿到参数列表
    t_code = request.GET.getlist('code', default=[])
    t_name = request.GET.get('name', default='-1')
    if t_code is [] or t_name is "":
        context = {
            'name_list': get_area_list(),
            'show_box': False,
        }
        return render(request, 'docxReport.html', context)

    v_list = models.Versions.objects.order_by("-id").all()[:7]
    vv_list = []
    vt_list = []
    for item in v_list:
        vv_list.append(item.id)
        vt_list.append((item.v_time + timedelta(hours=8)).strftime("%m-%d %H"))
    vt_list.reverse()
    areas_infos = []
    for i_code in t_code:
        areas_infos.append(models.AreaInfo.objects.filter(v_id__in=vv_list, code=i_code).all())

    document: Doc = Document()
    document.add_heading(t_name, 1)
    document.add_heading(text='Attractions', level=2)
    paragragh1 = document.add_paragraph()
    text1 = paragragh1.add_run('Attraction Name List')
    text1.font.size = Pt(11)
    text1.font.name = u'微软雅黑'
    paragragh1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = document.add_table(rows=len(areas_infos) + 1, cols=7)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table.style = 'Light List Accent 2'
    table.cell(0, 0).text = 'No.'
    table.cell(0, 1).text = 'Code'
    table.cell(0, 2).text = 'Name'
    table.cell(0, 3).text = 'Address'
    table.cell(0, 4).text = 'Grade'
    table.cell(0, 5).text = 'Business\nHours'
    table.cell(0, 6).text = 'Recent\nFlow'
    for row_index in range(len(areas_infos)):
        table.cell(row_index + 1, 0).text = str(row_index + 1)
        table.cell(row_index + 1, 1).text = areas_infos[row_index].first().code
        table.cell(row_index + 1, 2).text = areas_infos[row_index].first().name
        table.cell(row_index + 1, 3).text = areas_infos[row_index].first().address
        table.cell(row_index + 1, 4).text = areas_infos[row_index].first().grade
        table.cell(row_index + 1, 5).text = areas_infos[row_index].first().t_time
        table.cell(row_index + 1, 6).text = str(areas_infos[row_index].first().num)
    for row_index in range(len(areas_infos) + 1):
        for col_index in range(7):
            cell = table.cell(row_index, col_index)
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    document.add_heading(text='Chart with single data', level=2)

    for item in areas_infos:
        paragragh = document.add_paragraph()
        text = paragragh.add_run('Visitor flow chart of {0}'.format(item.first().name))
        text.font.size = Pt(11)
        text.font.name = u'微软雅黑'
        paragragh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        url = create_image.create_single(vt_list, item)
        paragragh = document.add_paragraph()
        paragragh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        text = paragragh.add_run()
        text.add_picture(image_path_or_stream=url, width=Inches(6.252))

    url = create_image.create_multi(vt_list, areas_infos)
    document.add_heading(text='Multi-attraction data comparison', level=2)
    paragragh = document.add_paragraph()
    text = paragragh.add_run('Chart with mutiple spot data')
    text.font.size = Pt(11)
    text.font.name = u'微软雅黑'
    paragragh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragragh = document.add_paragraph()
    paragragh.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    text = paragragh.add_run()
    text.add_picture(image_path_or_stream=url, width=Inches(6.252))

    file_name = './static/reports/' + t_name + '.docx'
    document.save(file_name)
    context = {
        'name_list': get_area_list(),
        'show_box': True,
        'docx_src': t_name,
    }
    return render(request, 'docxReport.html', context)


def read_file(file_name, chunk_size=512):
    with open(file_name, "rb") as f:
        while True:
            c = f.read(chunk_size)
            if c:
                yield c
            else:
                break


def download_docx_report(request):
    t_name = request.GET.get('fileName', default='-1')
    if t_name is '-1':
        return 404

    file_path = './static/reports/' + t_name + '.docx'
    response = StreamingHttpResponse(read_file(file_path))
    response["Content-Type"] = "application/octet-stream"
    response["Content-Disposition"] = 'attachment; filename={0}.docx'.format(urlquote(t_name))
    response["Access-Control-Expose-Headers"] = "Content-Disposition"  # 为了使前端获取到Content-Disposition属性

    return response


def send_email(request):
    email_names = request.POST['name']
    file_names = request.POST.getlist('code')
    names = email_names.split(';')
    try:
        email.send(names, file_names)
        resp = {"code": 200, "msg": "发送完毕"}
    except:
        resp = {"code": 500, "msg": "发送失败"}
    response = HttpResponse(json.dumps(resp), content_type="application/json")
    return response
